# -*- coding: utf-8 -*-
"""
Движок «Документы по встрече»: заполнение шаблонов и сборка документов.

Два режима (по видению владельца):

  Режим A (жёсткий) — есть DOCX-шаблон с плейсхолдерами {{поле}}. Система
    извлекает значения полей ИЗ ВСТРЕЧИ и заполняет шаблон, помечая КАЖДОЕ
    поле уверенностью: high (уверен) / low (не уверен, проверьте) / missing
    (не нашёл — оставил пустым). Плюс реквизиты компании (auto-слоты) и
    точный расчёт денег КОДОМ по позициям (модель числа не выдумывает).

  Режим B (гибкий) — есть пример/описание того, как мы делаем документ
    (напр. КП). Система сама формирует содержимое по встрече в стиле примера
    и рендерит в DOCX.

Этот модуль — чистое ядро (без БД/HTTP): извлечение плейсхолдеров, заполнение
DOCX с сохранением форматирования, расчёт денег, сумма прописью. LLM-часть
(извлечение полей по встрече, сборка контента) — тонкие обёртки, роут
прокидывает готовый llm_router и данные встречи.
"""
from __future__ import annotations

import io
import logging
import re
from typing import Any, Dict, List, Optional, Tuple

logger = logging.getLogger(__name__)

# {{ поле }} или {{ поле.вложенное }} — как в фирменных бланках Word.
_PLACEHOLDER_RE = re.compile(r"\{\{\s*([\w.\-]+)\s*\}\}")

# Реквизиты компании — автослоты (заполняются из профиля, не из встречи).
COMPANY_SLOT_PREFIXES = ("company", "supplier", "seller", "our")


# ── Плейсхолдеры ────────────────────────────────────────────────────────────

def extract_placeholders(docx_bytes: bytes) -> List[str]:
    """Список уникальных плейсхолдеров {{поле}} из DOCX (параграфы + таблицы),
    в порядке первого появления."""
    from docx import Document
    doc = Document(io.BytesIO(docx_bytes))
    seen: List[str] = []
    for text in _iter_block_texts(doc):
        for m in _PLACEHOLDER_RE.finditer(text):
            key = m.group(1)
            if key not in seen:
                seen.append(key)
    return seen


def _iter_block_texts(doc) -> List[str]:
    """Тексты всех параграфов документа, включая ячейки таблиц."""
    texts: List[str] = []
    for p in doc.paragraphs:
        texts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    texts.append(p.text)
    return texts


def _iter_all_paragraphs(doc):
    """Все параграфы (тело + таблицы) — для заполнения."""
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def _fill_paragraph(paragraph, values: Dict[str, str]) -> int:
    """Заменить {{поле}} в параграфе. Плейсхолдер может быть разбит между
    несколькими run'ами (Word так часто делает) — поэтому склеиваем текст
    параграфа, подставляем и кладём результат в ПЕРВЫЙ run, очищая остальные
    (сохраняет форматирование первого run'а плейсхолдера). Возвращает число
    сделанных подстановок."""
    full = "".join(run.text for run in paragraph.runs)
    if "{{" not in full:
        return 0

    count = [0]

    def _sub(m):
        key = m.group(1)
        if key in values:
            count[0] += 1
            return str(values[key])
        return m.group(0)  # неизвестный плейсхолдер оставляем как есть

    new = _PLACEHOLDER_RE.sub(_sub, full)
    if new == full:
        return 0
    if paragraph.runs:
        paragraph.runs[0].text = new
        for run in paragraph.runs[1:]:
            run.text = ""
    return count[0]


def fill_template(docx_bytes: bytes, values: Dict[str, Any],
                  table_items: Optional[List[Dict[str, Any]]] = None,
                  item_prefix: str = "items") -> bytes:
    """Заполнить DOCX-шаблон значениями {{поле}} → value. Возвращает bytes.

    table_items: список позиций для ПОСТРОЧНОГО заполнения таблицы — строка
    шаблона с плейсхолдерами {{items.name}}/{{items.qty}}/… размножается по
    каждой позиции (табличная часть договора/КП/счёта)."""
    from docx import Document
    doc = Document(io.BytesIO(docx_bytes))
    if table_items:
        _expand_table_rows(doc, table_items, item_prefix)
    str_values = {k: ("" if v is None else str(v)) for k, v in (values or {}).items()}
    filled = 0
    for p in _iter_all_paragraphs(doc):
        filled += _fill_paragraph(p, str_values)
    logger.info("[MeetingDoc] template filled: %d substitutions", filled)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _expand_table_rows(doc, items: List[Dict[str, Any]], prefix: str) -> None:
    """Найти в таблицах строку-шаблон с {{prefix.поле}} и размножить её по
    каждой позиции, заполнив значениями. Порядковый номер — {{prefix.no}}."""
    import copy
    marker = f"{{{{{prefix}."
    for table in doc.tables:
        tmpl_row = None
        row_idx = None  # индекс СРЕДИ СТРОК (не в parent-XML)
        for i, row in enumerate(table.rows):
            if any(marker in c.text for c in row.cells):
                tmpl_row = row
                row_idx = i
                break
        if tmpl_row is None:
            continue
        tmpl_tr = tmpl_row._tr
        parent = tmpl_tr.getparent()
        pos = list(parent).index(tmpl_tr)  # XML-позиция шаблонной строки
        # вставляем клоны ПЕРЕД шаблоном (порядок сохраняется), шаблон удаляем
        for idx in range(len(items)):
            new_tr = copy.deepcopy(tmpl_tr)
            parent.insert(pos + idx, new_tr)
        parent.remove(tmpl_tr)
        # добавленные строки теперь на row_idx .. row_idx+len-1
        rows = table.rows
        for idx, item in enumerate(items):
            row = rows[row_idx + idx]
            vals = {f"{prefix}.{k}": item.get(k) for k in item.keys()}
            vals[f"{prefix}.no"] = idx + 1
            flat = {k: ("" if v is None else str(v)) for k, v in vals.items()}
            for cell in row.cells:
                for p in cell.paragraphs:
                    _fill_paragraph(p, flat)


# ── Уверенность заполнения ──────────────────────────────────────────────────

# Значения уверенности, которые возвращает извлечение по встрече.
CONF_HIGH = "high"      # 🟢 уверен (нашёл прямо во встрече)
CONF_LOW = "low"        # 🟡 не уверен / предположил — проверьте
CONF_MISSING = "missing"  # 🔴 не нашёл — поле пустое, дозаполните


def build_confidence_report(fields: Dict[str, Dict[str, Any]],
                            company_keys: Optional[List[str]] = None) -> Dict[str, Any]:
    """Из карты полей {key: {value, confidence, quote}} собрать сводку
    уверенности для пользователя: что уверенно, что проверить, что пусто."""
    company_keys = set(company_keys or [])
    high, low, missing, from_company = [], [], [], []
    for key, f in (fields or {}).items():
        if key in company_keys:
            from_company.append(key)
            continue
        conf = (f or {}).get("confidence", CONF_MISSING)
        entry = {"key": key, "value": (f or {}).get("value", ""),
                 "quote": (f or {}).get("quote", "")}
        if conf == CONF_HIGH:
            high.append(entry)
        elif conf == CONF_LOW:
            low.append(entry)
        else:
            missing.append(entry)
    return {
        "confident": high,          # 🟢 подставил и уверен
        "uncertain": low,           # 🟡 подставил, но проверьте
        "unfilled": missing,        # 🔴 не заполнил
        "from_requisites": from_company,  # из реквизитов компании
        "counts": {"confident": len(high), "uncertain": len(low),
                   "unfilled": len(missing), "from_requisites": len(from_company)},
    }


def values_from_fields(fields: Dict[str, Dict[str, Any]]) -> Dict[str, str]:
    """Плоская карта key→value для заполнения (пустая строка для missing)."""
    out: Dict[str, str] = {}
    for key, f in (fields or {}).items():
        v = (f or {}).get("value")
        out[key] = "" if v is None else str(v)
    return out


# ── Деньги: расчёт КОДОМ (модель числа не выдумывает) ───────────────────────

def compute_line_items(items: List[Dict[str, Any]], vat_rate: float = 20.0) -> Dict[str, Any]:
    """Посчитать позиции сметы/КП/счёта ТОЧНО в коде.

    Адаптер над общим Decimal-движком (backend.core.documents.money:
    HALF_UP, цена квантуется до умножения, спецификация MeetFlow-тестов) —
    раньше здесь был float-дубль с расхождениями в копейках. Форма ответа
    сохранена (float'ы, ключи lines/subtotal/vat/total/total_in_words).
    """
    from backend.core.documents.money import compute_totals
    r = compute_totals(items or [], vat_rate=vat_rate, vat_included=False)
    return {
        "lines": [{
            "name": ln["name"], "unit": ln["unit"],
            "qty": float(ln["qty"]), "price": float(ln["price"]),
            "amount": float(ln["line_total"]),
        } for ln in r["lines"]],
        "subtotal": float(r["subtotal"]),
        "vat_rate": float(vat_rate or 0),
        "vat": float(r["vat_amount"]),
        "total": float(r["total"]),
        "total_in_words": r["total_in_words"],
    }


def _num(v: Any, default: float) -> float:
    try:
        if isinstance(v, str):
            v = v.replace(" ", "").replace(" ", "").replace(",", ".")
        return float(v)
    except (TypeError, ValueError):
        return default


# Сумма прописью: единая реализация в backend.core.documents.money
from backend.core.documents.money import rubles_in_words  # noqa: E402,F401

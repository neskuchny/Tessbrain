# -*- coding: utf-8 -*-
"""Заполнение документов по шаблону из встреч (перенос идеи MeetFlow).

«Прошла встреча → выбираем шаблон (договор/КП/карточка) → система
заполняет поля из встречи, реквизиты из профиля, деньги считает КОД,
и честно помечает каждое поле: найдено / предположено / не найдено».

Два режима:
- strict:   шаблон с плейсхолдерами {{поле}} — заполняются ТОЛЬКО поля,
            текст шаблона не меняется;
- flexible: шаблон = эталон стиля (пример КП) — система пишет документ
            по мотивам, факты из встреч, структура/тон из примера.

Хранилища per-user (tenant-паттерн): шаблоны и реквизиты — JSON.
"""
from __future__ import annotations

import json
import logging
import re
import uuid
from pathlib import Path
from typing import Any, Dict, List, Optional

from backend.core.documents.money import compute_totals

logger = logging.getLogger(__name__)

_PLACEHOLDER_RE = re.compile(r"\{\{\s*([^{}]+?)\s*\}\}")

# Плейсхолдеры денег: заполняются расчётом, не LLM
_MONEY_FIELDS = {
    "итого": "total", "сумма": "total", "сумма_итого": "total",
    "сумма_прописью": "total_in_words", "ндс": "vat_amount",
    "сумма_без_ндс": "subtotal", "подытог": "subtotal",
}

CONF_FOUND = "found"          # ✅ найдено во встрече/реквизитах
CONF_ASSUMED = "assumed"      # ⚠️ предположено (данных мало)
CONF_MISSING = "missing"      # ❌ не найдено — дозаполнить вручную


def extract_placeholders(template_text: str) -> List[str]:
    """Уникальные {{плейсхолдеры}} в порядке появления."""
    seen: List[str] = []
    for m in _PLACEHOLDER_RE.finditer(template_text or ""):
        name = m.group(1).strip()
        if name and name not in seen:
            seen.append(name)
    return seen


# ── per-user хранилища ──────────────────────────────────────────────────

def _store_dir(user_id: str) -> Path:
    from backend.core.store.tenant_paths import _DATA_ROOT, _require_uuid
    user_id = _require_uuid(user_id, "user_id")
    d = Path(_DATA_ROOT) / "documents" / user_id
    d.mkdir(parents=True, exist_ok=True)
    return d


def _read_json(path: Path, default):
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except (OSError, ValueError):
        return default


def _write_json(path: Path, data) -> None:
    from backend.core.store.tenant_io import atomic_write_json
    atomic_write_json(str(path), data)


def get_requisites(user_id: str) -> Dict[str, str]:
    """Реквизиты компании: {название, ИНН, банк, счёт, подписант, ...} —
    заполняются один раз, подставляются во все документы."""
    return _read_json(_store_dir(user_id) / "requisites.json", {})


def save_requisites(user_id: str, requisites: Dict[str, str]) -> Dict[str, str]:
    clean = {str(k).strip(): str(v).strip()
             for k, v in (requisites or {}).items() if str(k).strip()}
    _write_json(_store_dir(user_id) / "requisites.json", clean)
    return clean


def list_doc_templates(user_id: str) -> List[dict]:
    return _read_json(_store_dir(user_id) / "doc_templates.json", [])


def save_doc_template(user_id: str, *, title: str, body: str = "",
                      mode: str = "strict",
                      template_id: Optional[str] = None,
                      docx_b64: Optional[str] = None) -> dict:
    """Сохранить шаблон (upsert по id). mode: strict|flexible.
    docx_b64 — фирменный бланк Word: {{плейсхолдеры}} извлекаются из
    docx (meeting_doc_engine), бинарь хранится рядом, kind='docx'."""
    items = list_doc_templates(user_id)
    rec = None
    if template_id:
        rec = next((t for t in items if t.get("id") == template_id), None)
    if rec is None:
        rec = {"id": template_id or uuid.uuid4().hex[:10]}
        items.append(rec)
    if docx_b64:
        import base64
        from backend.core.documents.meeting_doc_engine import (
            extract_placeholders as docx_placeholders,
        )
        blob = base64.b64decode(docx_b64)
        (_store_dir(user_id) / f"tpl_{rec['id']}.docx").write_bytes(blob)
        rec.update({
            "title": (title or "Бланк Word").strip()[:120],
            "body": "", "kind": "docx", "mode": "strict",
            "placeholders": docx_placeholders(blob),
        })
    else:
        rec.update({
            "title": (title or "Шаблон").strip()[:120],
            "body": body or "",
            "kind": "text",
            "mode": mode if mode in ("strict", "flexible") else "strict",
            "placeholders": extract_placeholders(body or ""),
        })
    _write_json(_store_dir(user_id) / "doc_templates.json", items)
    return rec


def load_template_docx(user_id: str, template_id: str) -> Optional[bytes]:
    p = _store_dir(user_id) / f"tpl_{template_id}.docx"
    try:
        return p.read_bytes()
    except OSError:
        return None


# ── Заготовки контекста: типовые блоки «услуга + цена» и т.п. ──────────

def list_context_presets(user_id: str) -> List[dict]:
    return _read_json(_store_dir(user_id) / "context_presets.json", [])


def save_context_preset(user_id: str, *, title: str, text: str,
                        preset_id: Optional[str] = None) -> dict:
    """Сохранить заготовку (upsert): постоянно вставляемый блок контекста
    (описание типовой услуги, прайс, условия) — при заполнении отмечается
    галочкой вместо повторного набора."""
    items = list_context_presets(user_id)
    rec = None
    if preset_id:
        rec = next((p for p in items if p.get("id") == preset_id), None)
    if rec is None:
        rec = {"id": preset_id or uuid.uuid4().hex[:10]}
        items.append(rec)
    rec.update({"title": (title or "Заготовка").strip()[:120],
                "text": (text or "").strip()[:8000]})
    _write_json(_store_dir(user_id) / "context_presets.json", items)
    return rec


def delete_context_preset(user_id: str, preset_id: str) -> bool:
    items = list_context_presets(user_id)
    kept = [p for p in items if p.get("id") != preset_id]
    if len(kept) == len(items):
        return False
    _write_json(_store_dir(user_id) / "context_presets.json", kept)
    return True


def presets_text(user_id: str, preset_ids: List[str]) -> str:
    """Текст выбранных заготовок для подмешивания в extra_context."""
    if not preset_ids:
        return ""
    by_id = {p.get("id"): p for p in list_context_presets(user_id)}
    parts = [f"[{by_id[i]['title']}]\n{by_id[i]['text']}"
             for i in preset_ids if i in by_id]
    return "\n\n".join(parts)


def delete_doc_template(user_id: str, template_id: str) -> bool:
    items = list_doc_templates(user_id)
    kept = [t for t in items if t.get("id") != template_id]
    if len(kept) == len(items):
        return False
    _write_json(_store_dir(user_id) / "doc_templates.json", kept)
    return True


# ── сбор контекста встреч ───────────────────────────────────────────────

async def _meetings_text(user_id: str, meeting_ids: List[str],
                         cap_chars: int = 60000) -> str:
    from backend.core.sima.tessent_provider import fetch_meeting_transcript
    parts: List[str] = []
    for mid in (meeting_ids or [])[:5]:
        try:
            data = await fetch_meeting_transcript(user_id, mid)
        except Exception as e:
            logger.warning(f"fill: transcript {mid} failed: {e}")
            data = None
        if not data:
            continue
        # fetch_meeting_transcript отдаёт текст в "content" (саммари+моменты+
        # транскрипт, markdown), а заголовок — в metadata.title/name. Раньше здесь
        # читались НЕсуществующие ключи "transcript"/"text"/"title" → текст встречи
        # НИКОГДА не доходил до LLM (заполнялись только реквизиты, поля встречи ❌).
        meta = data.get("metadata") if isinstance(data.get("metadata"), dict) else {}
        title = (meta.get("title") or data.get("title") or data.get("name") or mid)
        txt = (data.get("content") or data.get("transcript") or data.get("text") or "")
        if not str(txt).strip():
            continue  # встреча без контента — не гоним пустой блок в LLM
        parts.append(f"=== ВСТРЕЧА: {title} ===\n{txt}")
    return "\n\n".join(parts)[:cap_chars]


def _req_match(field: str, requisites: Dict[str, str]) -> Optional[str]:
    fl = field.lower().replace("_", " ").strip()
    for k, v in (requisites or {}).items():
        kl = k.lower().replace("_", " ").strip()
        if kl == fl or kl in fl or fl in kl:
            return v
    return None


# ── главный вход ────────────────────────────────────────────────────────

async def fill_document(user_id: str, *, template_text: str = "",
                        meeting_ids: List[str],
                        extra_context: str = "",
                        mode: str = "strict",
                        vat_rate: float = 0,
                        vat_included: bool = False,
                        template_docx: Optional[bytes] = None,
                        llm=None) -> dict:
    """Заполнить документ из встреч. Возвращает:
    {document, fields:{имя:{value,confidence,source}}, items, money,
     unfilled:[...], mode, docx_b64?}.

    template_docx — фирменный бланк Word: {{поля}} заполняются прямо в
    docx (meeting_doc_engine, форматирование сохраняется), позиции
    размножаются по строке-шаблону {{items.*}}."""
    if llm is None:
        from backend.core.llm.router import get_llm_router
        llm = get_llm_router()

    requisites = get_requisites(user_id)
    meetings = await _meetings_text(user_id, meeting_ids)
    # Честно: выбраны встречи, но контента у них нет (нет транскрипта/саммари) и
    # доп-контекста тоже нет → раньше выходил «документ из одних ❌», и это читалось
    # как «ничего не заполнилось». Говорим прямо, что подставлять нечего.
    if meeting_ids and not meetings.strip() and not (extra_context or "").strip():
        raise ValueError(
            "У выбранных встреч нет транскрипта/саммари — подставить в бланк нечего. "
            "Выберите встречу с транскриптом или добавьте текст в «Доп. контекст».")
    context = "\n\n".join(x for x in (
        meetings,
        f"=== ДОП. КОНТЕКСТ ===\n{extra_context}" if extra_context else "",
    ) if x)

    if template_docx:
        return await _fill_docx(template_docx, context, requisites, llm,
                                vat_rate=vat_rate, vat_included=vat_included)
    if mode == "flexible":
        return await _fill_flexible(template_text, context, requisites, llm)
    return await _fill_strict(template_text, context, requisites, llm,
                              vat_rate=vat_rate, vat_included=vat_included)


async def _resolve_fields(placeholders: List[str], context: str,
                          requisites: Dict[str, str], llm, *,
                          vat_rate: float, vat_included: bool) -> tuple:
    """Общий резолв полей для текстового И docx-путей: реквизиты без LLM →
    остальное LLM строго JSON → деньги детерминированно по позициям.
    Возвращает (fields, items, money)."""
    fields: Dict[str, dict] = {}
    to_llm: List[str] = []
    for f in placeholders:
        key = f.lower().replace(" ", "_")
        if key in _MONEY_FIELDS or key.startswith("items."):
            continue                       # деньги/позиции посчитает код
        rv = _req_match(f, requisites)
        if rv:
            fields[f] = {"value": rv, "confidence": CONF_FOUND,
                         "source": "реквизиты компании"}
        else:
            to_llm.append(f)

    items: List[dict] = []
    if to_llm or placeholders:
        prompt = (
            "Заполни поля документа ФАКТАМИ из материалов встречи. "
            "НЕ выдумывай: если данных нет — верни null; если данных мало "
            "и ты ПРЕДПОЛАГАЕШЬ — пометь assumed.\n\n"
            f"МАТЕРИАЛЫ:\n{context[:45000]}\n\n"
            f"ПОЛЯ: {json.dumps(to_llm, ensure_ascii=False)}\n\n"
            "Если в материалах есть товары/услуги с ценами — верни их в items "
            "(суммы НЕ считай, посчитает код).\n"
            'Ответь ТОЛЬКО JSON: {"fields": {"имя_поля": {"value": "...", '
            '"confidence": "found|assumed", "source": "короткая цитата"}}, '
            '"items": [{"name": "...", "qty": 1, "unit": "шт", "price": 0}]}')
        try:
            data = await llm.generate_json(prompt=prompt, temperature=0.1)
        except Exception as e:
            logger.warning(f"fill resolve LLM failed: {e}")
            data = {}
        raw_fields = (data or {}).get("fields") or {}
        for f in to_llm:
            rf = raw_fields.get(f) or {}
            val = rf.get("value")
            if val in (None, "", "null"):
                fields[f] = {"value": None, "confidence": CONF_MISSING,
                             "source": None}
            else:
                conf = rf.get("confidence")
                fields[f] = {
                    "value": str(val),
                    "confidence": (CONF_ASSUMED if conf == "assumed"
                                   else CONF_FOUND),
                    "source": (rf.get("source") or "")[:200] or None}
        items = [i for i in ((data or {}).get("items") or [])
                 if isinstance(i, dict) and i.get("name")][:30]

    money = compute_totals(items, vat_rate=vat_rate,
                           vat_included=vat_included) if items else None
    for f in placeholders:
        key = f.lower().replace(" ", "_")
        if key in _MONEY_FIELDS:
            if money:
                fields[f] = {"value": str(money[_MONEY_FIELDS[key]]),
                             "confidence": CONF_FOUND,
                             "source": "расчёт по позициям (код)"}
            else:
                fields[f] = {"value": None, "confidence": CONF_MISSING,
                             "source": None}
    return fields, items, money


async def _fill_strict(template_text: str, context: str,
                       requisites: Dict[str, str], llm, *,
                       vat_rate: float, vat_included: bool) -> dict:
    placeholders = extract_placeholders(template_text)
    fields, items, money = await _resolve_fields(
        placeholders, context, requisites, llm,
        vat_rate=vat_rate, vat_included=vat_included)

    def _sub(m):
        name = m.group(1).strip()
        v = (fields.get(name) or {}).get("value")
        return str(v) if v not in (None, "") else m.group(0)
    document = _PLACEHOLDER_RE.sub(_sub, template_text)

    unfilled = [f for f in placeholders
                if (fields.get(f) or {}).get("confidence") == CONF_MISSING]
    return {"mode": "strict", "document": document, "fields": fields,
            "items": items,
            "money": _money_jsonable(money),
            "unfilled": unfilled, "placeholders": placeholders}


async def _fill_docx(template_docx: bytes, context: str,
                     requisites: Dict[str, str], llm, *,
                     vat_rate: float, vat_included: bool) -> dict:
    """Фирменный бланк Word: тот же резолв полей, рендер — прямо в docx
    (форматирование сохраняется, строка {{items.*}} размножается)."""
    import base64
    from backend.core.documents import meeting_doc_engine as E
    placeholders = E.extract_placeholders(template_docx)
    fields, items, money = await _resolve_fields(
        placeholders, context, requisites, llm,
        vat_rate=vat_rate, vat_included=vat_included)
    values = {k: (f.get("value") or "") for k, f in fields.items()}
    table_items = None
    if money:
        table_items = [{"name": ln["name"], "unit": ln["unit"],
                        "qty": float(ln["qty"]), "price": float(ln["price"]),
                        "amount": float(ln["line_total"])}
                       for ln in money["lines"]]
    blob = E.fill_template(template_docx, values, table_items=table_items)
    unfilled = [f for f in placeholders
                if not f.lower().startswith("items.")
                and (fields.get(f) or {}).get("confidence") == CONF_MISSING]
    preview = "\n".join(
        f"{k}: {v}" for k, v in values.items() if v) or "(поля не заполнены)"
    return {"mode": "docx", "document": preview, "fields": fields,
            "items": items, "money": _money_jsonable(money),
            "unfilled": unfilled, "placeholders": placeholders,
            "docx_b64": base64.b64encode(blob).decode()}


async def _fill_flexible(template_text: str, context: str,
                         requisites: Dict[str, str], llm) -> dict:
    """Шаблон = эталон стиля: структура/тон из примера, факты из встреч."""
    req_txt = "\n".join(f"- {k}: {v}" for k, v in (requisites or {}).items())
    prompt = (
        "Напиши документ ПО ОБРАЗЦУ: повтори структуру, тон и уровень "
        "детализации примера, но наполни ФАКТАМИ из материалов встречи. "
        "НЕ выдумывай факты; чего нет в материалах — пометь в тексте "
        "[уточнить: что именно]. Реквизиты компании используй как есть.\n\n"
        f"ОБРАЗЕЦ СТИЛЯ:\n{template_text[:15000]}\n\n"
        f"РЕКВИЗИТЫ:\n{req_txt or '—'}\n\n"
        f"МАТЕРИАЛЫ ВСТРЕЧ:\n{context[:45000]}\n\n"
        "Ответь ТОЛЬКО текстом документа (markdown).")
    try:
        resp = await llm.generate(prompt=prompt, temperature=0.3,
                                  max_tokens=4000)
        text = resp.get("text") if isinstance(resp, dict) else str(resp)
    except Exception as e:
        logger.warning(f"fill flexible LLM failed: {e}")
        text = ""
    unfilled = re.findall(r"\[уточнить:([^\]]*)\]", text or "")
    return {"mode": "flexible", "document": text or "",
            "fields": {}, "items": [], "money": None,
            "unfilled": [u.strip() for u in unfilled],
            "placeholders": []}


def _money_jsonable(money: Optional[dict]) -> Optional[dict]:
    if not money:
        return None
    out = dict(money)
    out["lines"] = [{k: str(v) for k, v in ln.items()}
                    for ln in money["lines"]]
    for k in ("subtotal", "vat_rate", "vat_amount", "total"):
        out[k] = str(out[k])
    return out


def build_example_template_docx() -> bytes:
    """«Рыба» — пример DOCX-шаблона с плейсхолдерами {{поле}}, который
    пользователь скачивает как отправную точку. Показывает и одиночные
    поля, и табличную часть {{items.*}} (сумму считает код, не модель)."""
    import io as _io

    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading("Коммерческое предложение", level=0)

    p = doc.add_paragraph()
    p.add_run("Клиент: ").bold = True
    p.add_run("{{client_name}}")
    p = doc.add_paragraph()
    p.add_run("Дата: ").bold = True
    p.add_run("{{date}}")
    p = doc.add_paragraph()
    p.add_run("Проект: ").bold = True
    p.add_run("{{project}}")

    doc.add_paragraph()
    doc.add_paragraph(
        "Настоящее предложение подготовлено для {{client_name}} "
        "в рамках проекта «{{project}}».")

    doc.add_heading("Состав работ / позиции", level=1)
    table = doc.add_table(rows=2, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Наименование", "Кол-во", "Цена"
    row = table.rows[1].cells
    # одна строка-образец: система размножит по числу позиций
    row[0].text, row[1].text, row[2].text = (
        "{{items.name}}", "{{items.qty}}", "{{items.price}}")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Итого: ").bold = True
    p.add_run("{{total}}").bold = True

    doc.add_paragraph()
    note = doc.add_paragraph(
        "Подсказка: замените плейсхолдеры {{...}} на свои поля и удалите этот "
        "абзац. Табличную строку {{items.*}} оставьте одну — система "
        "размножит её по числу позиций и посчитает {{total}} сама.")
    for r in note.runs:
        r.font.size = Pt(8)
        r.italic = True

    buf = _io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

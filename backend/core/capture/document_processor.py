# -*- coding: utf-8 -*-
"""
Document Processor - обработка и индексация документов.

Поддерживает:
- PDF файлы
- DOCX файлы
- TXT файлы
- Markdown файлы

Основано на подходе KAG для извлечения знаний из документов.

Интеграция:
- SemanticSplitter для умного разбиения текста
- KnowledgeExtractor для извлечения структурированных знаний
"""
import hashlib
import json
import logging
import re
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

# Импортируем модуль чанкинга
from ..chunking import KnowledgeExtractor, SemanticSplitter, SplitterConfig

logger = logging.getLogger(__name__)


@dataclass
class DocumentChunk:
    """Чанк документа для индексации."""
    id: str
    document_id: str
    content: str
    chunk_index: int
    start_char: int
    end_char: int
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class ProcessedDocument:
    """Обработанный документ."""
    id: str
    filename: str
    file_type: str
    title: str
    content: str
    chunks: List[DocumentChunk]
    metadata: Dict[str, Any]
    entities: List[Dict[str, Any]]
    processed_at: str


class DocumentProcessor:
    """
    Процессор документов для извлечения и индексации знаний.

    Реализует подход KAG:
    1. Извлечение текста из документов
    2. Чанкинг с перекрытием
    3. Извлечение сущностей с помощью LLM
    4. Индексация в граф и векторное хранилище
    """

    # Параметры чанкинга
    DEFAULT_CHUNK_SIZE = 1000  # символов
    DEFAULT_CHUNK_OVERLAP = 200  # символов

    # Поддерживаемые форматы
    SUPPORTED_FORMATS = {
        ".pdf": "pdf",
        ".docx": "docx",
        ".doc": "doc",
        ".txt": "text",
        ".md": "markdown",
        ".markdown": "markdown",
    }

    def __init__(
        self,
        graph_builder=None,
        vector_indexer=None,
        llm_client=None,
        chunk_size: int = DEFAULT_CHUNK_SIZE,
        chunk_overlap: int = DEFAULT_CHUNK_OVERLAP,
        use_semantic_chunking: bool = True
    ):
        """
        Args:
            graph_builder: GraphBuilder для добавления в граф
            vector_indexer: VectorIndexer для индексации
            llm_client: LLM клиент для извлечения сущностей
            chunk_size: Размер чанка в символах
            chunk_overlap: Перекрытие между чанками
            use_semantic_chunking: Использовать семантическое разбиение (NEW!)
        """
        self.graph_builder = graph_builder
        self.vector_indexer = vector_indexer
        self.llm_client = llm_client
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.use_semantic_chunking = use_semantic_chunking

        # Инициализируем семантический сплиттер
        self.semantic_splitter = SemanticSplitter(
            config=SplitterConfig(
                chunk_size=chunk_size,
                window_size=chunk_overlap,
                min_chunk_size=50
            ),
            llm_router=llm_client
        )

        # Инициализируем экстрактор знаний
        self.knowledge_extractor = KnowledgeExtractor(llm_router=llm_client)

        # Проверяем доступность библиотек
        self._check_dependencies()

        logger.info(f"✅ DocumentProcessor initialized (semantic_chunking={use_semantic_chunking})")

    def _check_dependencies(self):
        """Проверить доступность необходимых библиотек."""
        self._has_pypdf = False
        self._has_docx = False

        try:
            import pypdf
            self._has_pypdf = True
        except ImportError:
            logger.warning("⚠️ pypdf не установлен. PDF файлы не будут обрабатываться.")

        try:
            import docx
            self._has_docx = True
        except ImportError:
            logger.warning("⚠️ python-docx не установлен. DOCX файлы не будут обрабатываться.")

    async def process_file(
        self,
        file_path: str,
        metadata: Optional[Dict[str, Any]] = None,
        extract_entities: bool = True
    ) -> ProcessedDocument:
        """
        Обработать файл документа.

        Args:
            file_path: Путь к файлу
            metadata: Дополнительные метаданные
            extract_entities: Извлекать ли сущности с помощью LLM

        Returns:
            ProcessedDocument с извлечённым содержимым
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"Файл не найден: {file_path}")

        ext = path.suffix.lower()
        if ext not in self.SUPPORTED_FORMATS:
            raise ValueError(f"Неподдерживаемый формат: {ext}")

        file_type = self.SUPPORTED_FORMATS[ext]

        logger.info(f"📄 Обработка документа: {path.name} ({file_type})")

        # Генерируем ID документа
        doc_id = self._generate_doc_id(file_path)

        # Извлекаем текст
        content = await self._extract_text(path, file_type)

        if not content or len(content.strip()) < 50:
            raise ValueError(f"Документ пуст или слишком короткий: {file_path}")

        # Определяем заголовок
        title = self._extract_title(content, path.stem)

        # Создаём чанки (используем семантическое разбиение если включено)
        if self.use_semantic_chunking:
            semantic_chunks = self.semantic_splitter.split(
                text=content,
                source_id=doc_id,
                method="auto"  # Автоматический выбор метода
            )
            # Конвертируем Chunk в DocumentChunk
            chunks = [
                DocumentChunk(
                    id=sc.id,
                    document_id=doc_id,
                    content=sc.content,
                    chunk_index=sc.position,
                    start_char=sc.start_char,
                    end_char=sc.end_char,
                    metadata={
                        "type": sc.type.value if hasattr(sc.type, 'value') else str(sc.type),
                        "topic": sc.topic,
                        "speaker": sc.speaker,
                        "char_count": sc.char_count,
                        "word_count": sc.word_count,
                        "content_hash": sc.content_hash
                    }
                )
                for sc in semantic_chunks
            ]
            logger.info(f"   📑 Создано {len(chunks)} семантических чанков")
        else:
            chunks = self._create_chunks(doc_id, content)
            logger.info(f"   📑 Создано {len(chunks)} чанков (простое разбиение)")

        # Извлекаем сущности
        entities = []
        if extract_entities and self.llm_client:
            entities = await self._extract_entities(content, title)
            logger.info(f"   🏷️ Извлечено {len(entities)} сущностей")

        # Извлекаем структурированные знания из чанков (NEW!)
        knowledge_items = []
        if extract_entities and self.llm_client and self.use_semantic_chunking:
            for chunk in chunks[:5]:  # Ограничиваем первыми 5 чанками для скорости
                try:
                    items = await self.knowledge_extractor.extract(
                        text=chunk.content,
                        source_id=doc_id,
                        chunk_id=chunk.id,
                        context=f"Документ: {title}"
                    )
                    knowledge_items.extend(items)
                except Exception as e:
                    logger.warning(f"⚠️ Knowledge extraction failed for chunk {chunk.id}: {e}")

            if knowledge_items:
                logger.info(f"   📚 Извлечено {len(knowledge_items)} элементов знаний")

        # Собираем метаданные
        doc_metadata = {
            "filename": path.name,
            "file_path": str(path.absolute()),
            "file_size": path.stat().st_size,
            "file_type": file_type,
            "created_at": datetime.fromtimestamp(path.stat().st_ctime).isoformat(),
            "modified_at": datetime.fromtimestamp(path.stat().st_mtime).isoformat(),
            "char_count": len(content),
            "word_count": len(content.split()),
            "chunk_count": len(chunks),
            "knowledge_items_count": len(knowledge_items),
            "semantic_chunking": self.use_semantic_chunking,
            **(metadata or {})
        }

        # Добавляем knowledge_items в entities для обратной совместимости
        for ki in knowledge_items:
            entities.append({
                "name": ki.title,
                "type": f"Knowledge_{ki.type.value if hasattr(ki.type, 'value') else str(ki.type)}",
                "description": ki.content[:200],
                "confidence": ki.confidence,
                "importance": ki.importance,
                "tags": ki.tags,
                "knowledge_id": ki.id
            })

        return ProcessedDocument(
            id=doc_id,
            filename=path.name,
            file_type=file_type,
            title=title,
            content=content,
            chunks=chunks,
            metadata=doc_metadata,
            entities=entities,
            processed_at=datetime.now().isoformat()
        )

    async def process_and_index(
        self,
        file_path: str,
        project_id: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Обработать документ и добавить в индексы.

        Args:
            file_path: Путь к файлу
            project_id: ID проекта для связи
            metadata: Дополнительные метаданные

        Returns:
            Результат индексации
        """
        # Обрабатываем документ
        doc = await self.process_file(file_path, metadata)

        result = {
            "document_id": doc.id,
            "filename": doc.filename,
            "chunks_indexed": 0,
            "entities_added": 0,
            "knowledge_items_added": 0,
            "graph_nodes_added": 0
        }

        # Добавляем в граф
        if self.graph_builder:
            # Добавляем узел документа
            doc_node_id = f"Document_{doc.id}"
            await self.graph_builder.create_node(
                node_id=doc_node_id,
                label="Document",
                properties={
                    "id": doc.id,
                    "name": doc.title,
                    "filename": doc.filename,
                    "file_type": doc.file_type,
                    "content_preview": doc.content[:500] if doc.content else "",
                    "processed_at": doc.processed_at
                }
            )
            result["graph_nodes_added"] += 1

            # Связываем с проектом
            if project_id:
                await self.graph_builder._create_relationship(
                    from_id=doc_node_id,
                    to_id=f"Project_{project_id[:8]}",
                    rel_type="BELONGS_TO",
                    properties={}
                )

            # Добавляем сущности
            for entity in doc.entities:
                entity_name = entity.get("name", "Unknown")
                entity_type = entity.get("type", "Unknown")

                # Определяем тип узла (Knowledge или Entity)
                if entity_type.startswith("Knowledge_"):
                    # Это элемент знания
                    knowledge_id = entity.get("knowledge_id", f"Knowledge_{entity_name.replace(' ', '_')[:20]}_{doc.id[:8]}")
                    await self.graph_builder.create_node(
                        node_id=knowledge_id,
                        label="Knowledge",
                        properties={
                            "name": entity_name,
                            "knowledge_type": entity_type.replace("Knowledge_", ""),
                            "description": entity.get("description", ""),
                            "confidence": entity.get("confidence", 0.8),
                            "importance": entity.get("importance", 0.5),
                            "tags": entity.get("tags", []),
                            "source_document": doc.id
                        }
                    )

                    # Связываем с документом
                    await self.graph_builder._create_relationship(
                        from_id=doc_node_id,
                        to_id=knowledge_id,
                        rel_type="CONTAINS_KNOWLEDGE",
                        properties={}
                    )
                    result["knowledge_items_added"] += 1
                else:
                    # Обычная сущность
                    entity_id = f"Entity_{entity_name.replace(' ', '_')[:20]}_{doc.id[:8]}"
                    await self.graph_builder.create_node(
                        node_id=entity_id,
                        label="Entity",
                        properties={
                            "name": entity_name,
                            "entity_type": entity_type,
                            "description": entity.get("description", ""),
                            "source_document": doc.id
                        }
                    )

                    # Связываем с документом
                    await self.graph_builder._create_relationship(
                        from_id=doc_node_id,
                        to_id=entity_id,
                        rel_type="MENTIONS",
                        properties={}
                    )
                    result["entities_added"] += 1

            result["graph_nodes_added"] += result["entities_added"] + result["knowledge_items_added"]

        # Индексируем чанки
        if self.vector_indexer:
            for chunk in doc.chunks:
                await self.vector_indexer.add_document(
                    text=chunk.content,
                    metadata={
                        "id": chunk.id,
                        "document_id": doc.id,
                        "filename": doc.filename,
                        "title": doc.title,
                        "chunk_index": chunk.chunk_index,
                        "label": "DocumentChunk",
                        **chunk.metadata
                    }
                )
                result["chunks_indexed"] += 1

        logger.info(f"✅ Документ проиндексирован: {doc.filename}")
        logger.info(f"   📊 Чанков: {result['chunks_indexed']}, Сущностей: {result['entities_added']}, Знаний: {result['knowledge_items_added']}")

        return result

    async def _extract_text(self, path: Path, file_type: str) -> str:
        """Извлечь текст из файла."""
        if file_type == "pdf":
            return await self._extract_pdf(path)
        elif file_type in ("docx", "doc"):
            return await self._extract_docx(path)
        elif file_type in ("text", "markdown"):
            return await self._extract_text_file(path)
        else:
            raise ValueError(f"Неизвестный тип файла: {file_type}")

    async def _extract_pdf(self, path: Path) -> str:
        """Извлечь текст из PDF."""
        if not self._has_pypdf:
            raise ImportError("pypdf не установлен. Установите: pip install pypdf")

        import pypdf

        text_parts = []

        with open(path, "rb") as f:
            reader = pypdf.PdfReader(f)

            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"--- Страница {page_num + 1} ---\n{page_text}")

        return "\n\n".join(text_parts)

    async def _extract_docx(self, path: Path) -> str:
        """Извлечь текст из DOCX."""
        if not self._has_docx:
            raise ImportError("python-docx не установлен. Установите: pip install python-docx")

        import docx

        doc = docx.Document(path)

        text_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Извлекаем текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        return "\n\n".join(text_parts)

    async def _extract_text_file(self, path: Path) -> str:
        """Извлечь текст из текстового файла."""
        encodings = ["utf-8", "cp1251", "latin-1"]

        for encoding in encodings:
            try:
                with open(path, "r", encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue

        raise ValueError(f"Не удалось определить кодировку файла: {path}")

    def _extract_title(self, content: str, default_name: str) -> str:
        """Извлечь заголовок из содержимого."""
        # Пробуем найти заголовок в первых строках
        lines = content.split("\n")[:10]

        for line in lines:
            line = line.strip()

            # Markdown заголовок
            if line.startswith("# "):
                return line[2:].strip()

            # Заголовок в верхнем регистре или с заглавной буквы (первая непустая строка)
            if len(line) > 5 and len(line) < 200:
                # Не содержит много пунктуации
                if sum(1 for c in line if c in ".,;:!?()[]{}") < len(line) * 0.1:
                    return line

        return default_name

    def _create_chunks(self, doc_id: str, content: str) -> List[DocumentChunk]:
        """Создать чанки из содержимого."""
        chunks = []

        # Нормализуем пробелы
        content = re.sub(r'\s+', ' ', content)

        start = 0
        chunk_index = 0

        while start < len(content):
            end = start + self.chunk_size

            # Ищем хорошую точку разрыва (конец предложения)
            if end < len(content):
                # Ищем конец предложения в последних 20% чанка
                search_start = start + int(self.chunk_size * 0.8)
                search_text = content[search_start:end]

                # Ищем точку, восклицательный или вопросительный знак
                for sep in [". ", "! ", "? ", ".\n", "!\n", "?\n"]:
                    last_sep = search_text.rfind(sep)
                    if last_sep != -1:
                        end = search_start + last_sep + len(sep)
                        break

            chunk_text = content[start:end].strip()

            if chunk_text:
                chunk_id = f"{doc_id}_chunk_{chunk_index}"

                chunks.append(DocumentChunk(
                    id=chunk_id,
                    document_id=doc_id,
                    content=chunk_text,
                    chunk_index=chunk_index,
                    start_char=start,
                    end_char=end,
                    metadata={
                        "char_count": len(chunk_text),
                        "word_count": len(chunk_text.split())
                    }
                ))

                chunk_index += 1

            # Следующий чанк с перекрытием
            start = end - self.chunk_overlap

            # Защита от бесконечного цикла
            if start >= len(content) - 10:
                break

        return chunks

    async def _extract_entities(self, content: str, title: str) -> List[Dict[str, Any]]:
        """Извлечь сущности из документа с помощью LLM."""
        if not self.llm_client:
            return []

        # Берём первые 4000 символов для анализа
        sample = content[:4000]

        prompt = f"""Проанализируй документ и извлеки ключевые сущности.

Документ: "{title}"

Содержание (фрагмент):
{sample}

Извлеки сущности следующих типов:
- Люди (Person): имена, должности
- Организации (Organization): компании, отделы
- Проекты (Project): названия проектов
- Технологии (Technology): инструменты, языки, платформы
- Продукты (Product): продукты, сервисы
- Концепции (Concept): ключевые термины, идеи

Ответь в формате JSON:
{{
  "entities": [
    {{"name": "...", "type": "Person|Organization|Project|Technology|Product|Concept", "description": "краткое описание"}},
    ...
  ]
}}

Только JSON, без пояснений. Максимум 15 сущностей."""

        try:
            response = await self.llm_client.generate(
                prompt=prompt,
                temperature=0.1,
                max_tokens=1500
            )

            text = response if isinstance(response, str) else response.get("text", "")

            # Парсим JSON
            json_match = re.search(r'\{[\s\S]*\}', text)
            if json_match:
                data = json.loads(json_match.group())
                return data.get("entities", [])

        except Exception as e:
            logger.warning(f"⚠️ Ошибка извлечения сущностей: {e}")

        return []

    def _generate_doc_id(self, file_path: str) -> str:
        """Сгенерировать уникальный ID документа."""
        path = Path(file_path)

        # Используем хеш от пути и времени модификации
        stat = path.stat()
        content = f"{path.absolute()}_{stat.st_mtime}_{stat.st_size}"

        hash_obj = hashlib.md5(content.encode())
        return f"doc_{hash_obj.hexdigest()[:12]}"

    async def process_directory(
        self,
        directory_path: str,
        recursive: bool = True,
        project_id: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Обработать все документы в директории.

        Args:
            directory_path: Путь к директории
            recursive: Обрабатывать поддиректории
            project_id: ID проекта для связи

        Returns:
            Статистика обработки
        """
        path = Path(directory_path)

        if not path.exists() or not path.is_dir():
            raise ValueError(f"Директория не найдена: {directory_path}")

        stats = {
            "total_files": 0,
            "processed": 0,
            "failed": 0,
            "skipped": 0,
            "total_chunks": 0,
            "total_entities": 0,
            "errors": []
        }

        # Собираем файлы
        if recursive:
            files = list(path.rglob("*"))
        else:
            files = list(path.glob("*"))

        for file_path in files:
            if not file_path.is_file():
                continue

            ext = file_path.suffix.lower()
            if ext not in self.SUPPORTED_FORMATS:
                stats["skipped"] += 1
                continue

            stats["total_files"] += 1

            try:
                result = await self.process_and_index(
                    str(file_path),
                    project_id=project_id
                )

                stats["processed"] += 1
                stats["total_chunks"] += result.get("chunks_indexed", 0)
                stats["total_entities"] += result.get("entities_added", 0)

            except Exception as e:
                stats["failed"] += 1
                stats["errors"].append({
                    "file": str(file_path),
                    "error": str(e)
                })
                logger.error(f"❌ Ошибка обработки {file_path}: {e}")

        logger.info("📁 Обработка директории завершена:")
        logger.info(f"   ✅ Обработано: {stats['processed']}/{stats['total_files']}")
        logger.info(f"   ❌ Ошибок: {stats['failed']}")
        logger.info(f"   📑 Чанков: {stats['total_chunks']}")
        logger.info(f"   🏷️ Сущностей: {stats['total_entities']}")

        return stats


class DocumentSearcher:
    """
    Поиск по документам.

    Использует векторный поиск для нахождения релевантных чанков
    и возвращает контекст из документов.
    """

    def __init__(self, vector_indexer=None, graph_builder=None):
        """
        Args:
            vector_indexer: VectorIndexer для поиска
            graph_builder: GraphBuilder для получения метаданных
        """
        self.vector_indexer = vector_indexer
        self.graph_builder = graph_builder

    async def search(
        self,
        query: str,
        limit: int = 10,
        min_score: float = 0.5
    ) -> List[Dict[str, Any]]:
        """
        Поиск по документам.

        Args:
            query: Поисковый запрос
            limit: Максимальное количество результатов
            min_score: Минимальный порог релевантности

        Returns:
            Список найденных чанков с контекстом
        """
        if not self.vector_indexer:
            return []

        results = await self.vector_indexer.search(
            query=query,
            collection="tessent_documents",
            limit=limit * 2  # Берём больше для фильтрации
        )

        # Фильтруем по score
        filtered = [r for r in results if r.get("score", 0) >= min_score]

        # Группируем по документам
        by_document = {}
        for r in filtered:
            doc_id = r.get("metadata", {}).get("document_id", "unknown")
            if doc_id not in by_document:
                by_document[doc_id] = []
            by_document[doc_id].append(r)

        # Формируем результаты
        output = []
        for doc_id, chunks in by_document.items():
            # Берём лучший чанк
            best_chunk = max(chunks, key=lambda x: x.get("score", 0))

            output.append({
                "document_id": doc_id,
                "filename": best_chunk.get("metadata", {}).get("filename", ""),
                "title": best_chunk.get("metadata", {}).get("title", ""),
                "content": best_chunk.get("text", ""),
                "score": best_chunk.get("score", 0),
                "chunk_index": best_chunk.get("metadata", {}).get("chunk_index", 0),
                "total_chunks_found": len(chunks)
            })

        # Сортируем по релевантности
        output.sort(key=lambda x: -x["score"])

        return output[:limit]

    async def get_document_context(
        self,
        document_id: str,
        chunk_index: int,
        context_chunks: int = 2
    ) -> Dict[str, Any]:
        """
        Получить расширенный контекст из документа.

        Args:
            document_id: ID документа
            chunk_index: Индекс чанка
            context_chunks: Количество соседних чанков для контекста

        Returns:
            Расширенный контекст
        """
        if not self.vector_indexer:
            return {}

        # Собираем соседние чанки
        chunks = []
        for i in range(chunk_index - context_chunks, chunk_index + context_chunks + 1):
            if i < 0:
                continue

            chunk_id = f"{document_id}_chunk_{i}"

            # Пробуем получить чанк
            # (это зависит от реализации vector_indexer)
            try:
                result = await self.vector_indexer.get_by_id(
                    collection="tessent_documents",
                    id=chunk_id
                )
                if result:
                    chunks.append({
                        "index": i,
                        "content": result.get("text", ""),
                        "is_target": i == chunk_index
                    })
            except Exception:
                logger.debug("suppressed exception", exc_info=True)

        # Собираем контекст
        context_text = "\n\n".join(
            f"[{'>>>' if c['is_target'] else '   '}] {c['content']}"
            for c in sorted(chunks, key=lambda x: x["index"])
        )

        return {
            "document_id": document_id,
            "target_chunk": chunk_index,
            "context_range": f"{chunk_index - context_chunks} - {chunk_index + context_chunks}",
            "context_text": context_text,
            "chunks_found": len(chunks)
        }


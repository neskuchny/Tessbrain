"""Тесты для Analysis Engine — docx export + agent_tool."""
from __future__ import annotations

import asyncio
import json
import sys
from pathlib import Path

import pytest

_ROOT = Path(__file__).resolve().parents[2]
if str(_ROOT) not in sys.path:
    sys.path.insert(0, str(_ROOT))

from backend.core.analysis import export as export_mod


def _run(coro):
    return asyncio.run(coro)


# =============================================================================
# docx export
# =============================================================================


def test_markdown_to_docx_returns_bytes():
    md = "# Заголовок\n\n## Раздел\n\n- пункт один\n- **жирный** пункт\n\nОбычный текст."
    data = export_mod.markdown_to_docx_bytes(md, title="Тест")
    assert isinstance(data, bytes)
    assert len(data) > 0
    # .docx — это ZIP (начинается с PK).
    assert data[:2] == b"PK"


def test_markdown_to_docx_is_valid_docx():
    """Сгенерированный docx читается обратно python-docx."""
    md = "# Отчёт\n## Долг\n- значение 150\n---\n## Риски\n- иск X"
    data = export_mod.markdown_to_docx_bytes(md)
    import io

    from docx import Document
    doc = Document(io.BytesIO(data))
    texts = [p.text for p in doc.paragraphs]
    joined = "\n".join(texts)
    assert "Отчёт" in joined
    assert "Долг" in joined
    assert "значение 150" in joined


def test_markdown_bold_parsed():
    data = export_mod.markdown_to_docx_bytes("Текст **важное** слово")
    import io

    from docx import Document
    doc = Document(io.BytesIO(data))
    # Найдём параграф с bold-run.
    has_bold = any(
        run.bold for p in doc.paragraphs for run in p.runs if run.bold
    )
    assert has_bold


def test_report_to_docx_bytes():
    report = {"markdown": "# Разбор\n## Вывод\n- всё ок"}
    data = export_mod.report_to_docx_bytes(report, playbook_name="DD")
    assert data[:2] == b"PK"


def test_report_to_docx_empty_report():
    data = export_mod.report_to_docx_bytes({}, playbook_name="X")
    assert data[:2] == b"PK"  # пустой, но валидный docx


def test_markdown_to_docx_empty_string():
    data = export_mod.markdown_to_docx_bytes("")
    assert data[:2] == b"PK"


# =============================================================================
# agent_tool
# =============================================================================


def test_tool_schema_well_formed():
    from backend.core.analysis.agent_tool import TOOL_SCHEMA
    assert TOOL_SCHEMA["name"] == "run_document_analysis"
    props = TOOL_SCHEMA["parameters"]["properties"]
    assert "playbook_id" in props
    assert "document_ids" in props
    assert "client_context" in props
    assert TOOL_SCHEMA["parameters"]["required"] == ["playbook_id", "document_ids"]


def test_tool_requires_args():
    from backend.core.analysis.agent_tool import run_document_analysis_tool
    out = _run(run_document_analysis_tool("", [], user_id="u"))
    parsed = json.loads(out)
    assert parsed["ok"] is False
    assert "required" in parsed["error"]


def test_tool_unknown_playbook(monkeypatch):
    """Playbook не найден → ok=False с понятной ошибкой."""
    from backend.core.analysis import agent_tool

    # Postgres недоступен → mem store (пустой) → playbook не найдётся.
    out = _run(agent_tool.run_document_analysis_tool(
        "apb_nope", ["d1"], user_id="u",
    ))
    parsed = json.loads(out)
    assert parsed["ok"] is False
    assert "not found" in parsed["error"]


def test_tool_full_run(monkeypatch):
    """Интеграция: создать playbook в mem store, прогнать tool с fake LLM."""
    from backend.core.analysis import agent_tool, run_engine, store
    from backend.core.analysis.models import (
        AnalysisDimension,
        AnalysisPlaybook,
    )

    # Общий mem-store: tool создаёт свои store с postgres=None, но каждый
    # инстанс — свой dict. Подменяем фабрики на синглтоны.
    shared_pb = store.PlaybookStore(postgres=None)
    shared_run = store.RunStore(postgres=None)

    # Положим playbook.
    pb = AnalysisPlaybook(
        id="apb_x", name="DD",
        dimensions=[AnalysisDimension(key="debt", label="Долг", kind="quantitative",
                                      computation="debt", inputs=["debt"])],
    )
    _run(shared_pb.create(pb))

    # Подменяем конструкторы store на возврат синглтонов.
    monkeypatch.setattr(store, "PlaybookStore", lambda **kw: shared_pb)
    monkeypatch.setattr(store, "RunStore", lambda **kw: shared_run)
    # agent_tool импортирует store-классы внутри функции из module store —
    # патчим там же.
    monkeypatch.setattr(agent_tool, "RunStore", shared_run, raising=False)

    # Fake document loader + LLM в run_engine deps.
    docs = {"d1": {"id": "d1", "title": "T", "content": "Кредиты 100 млн"}}

    class _LLM:
        async def generate_json(self, prompt, system_prompt=None):
            return {"debt": [{"value": 100, "quote": "Кредиты 100 млн", "unit": "млн"}]}

    async def _loader(doc_id):
        return docs.get(doc_id)

    async def _fake_default_deps():
        async def _save(r):
            await shared_run.save(r)
        return run_engine.RunDeps(document_loader=_loader, llm=_LLM(), save_run=_save)

    monkeypatch.setattr(run_engine, "_default_deps", _fake_default_deps)

    # dispatch_run внутри создаёт свои store — подменим get_postgres на None
    # и store-классы (уже сделано). Но dispatch создаёт RunStore(postgres=pg)
    # напрямую из store module — синглтон вернётся.
    out = _run(agent_tool.run_document_analysis_tool(
        "apb_x", ["d1"], user_id="u", client_context="сделка",
    ))
    parsed = json.loads(out)
    assert parsed["ok"] is True
    assert parsed["status"] == "done"
    assert parsed["run_id"].startswith("arun_")
    assert "report_url" in parsed

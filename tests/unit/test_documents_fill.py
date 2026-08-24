# -*- coding: utf-8 -*-
"""Заполнение документов по шаблону (MeetFlow-перенос): деньги считает
код (Decimal/HALF_UP/прописью), поля помечаются found/assumed/missing,
реквизиты подставляются без LLM, оверрайды/шаблоны — per-user store."""
from __future__ import annotations

import asyncio
import json
from decimal import Decimal

import pytest

import backend.core.documents.fill_engine as fe
from backend.core.documents.fill_engine import (
    CONF_FOUND,
    CONF_MISSING,
    extract_placeholders,
    save_doc_template,
)
from backend.core.documents.money import compute_totals, rubles_in_words


# ── money: спецификация из MeetFlow-тестов ──────────────────────────────

def test_vat_on_top():
    r = compute_totals([{"name": "Работа", "qty": 2, "unit": "шт",
                         "price": 100}], vat_rate=20, vat_included=False)
    assert r["subtotal"] == Decimal("200.00")
    assert r["vat_amount"] == Decimal("40.00")
    assert r["total"] == Decimal("240.00")
    assert r["total_in_words"] == "Двести сорок рублей 00 копеек"


def test_vat_included():
    r = compute_totals([{"name": "Работа", "qty": 1, "price": 120}],
                       vat_rate=20, vat_included=True)
    assert r["vat_amount"] == Decimal("20.00")
    assert r["total"] == Decimal("120.00")


def test_rounding_and_quantize():
    r = compute_totals([{"name": "x", "qty": 1.5, "price": 99.99}], vat_rate=0)
    assert r["lines"][0]["line_total"] == Decimal("149.99")
    r2 = compute_totals([{"name": "x", "qty": 2, "price": 10.999}], vat_rate=0)
    assert r2["lines"][0]["price"] == Decimal("11.00")
    assert r2["subtotal"] == Decimal("22.00")


def test_defaults_and_empty():
    r = compute_totals([{"name": "x", "unit": "", "price": None}], vat_rate=0)
    assert r["lines"][0]["qty"] == Decimal("1")
    assert r["lines"][0]["unit"] == "шт"
    assert compute_totals([], vat_rate=20)["total_in_words"] == \
        "Ноль рублей 00 копеек"


def test_words_declensions():
    assert rubles_in_words(1) == "Один рубль 00 копеек"
    assert rubles_in_words(2) == "Два рубля 00 копеек"
    assert rubles_in_words(12) == "Двенадцать рублей 00 копеек"
    assert rubles_in_words("1234.50") == \
        "Одна тысяча двести тридцать четыре рубля 50 копеек"
    assert rubles_in_words(2000000) == "Два миллиона рублей 00 копеек"


# ── плейсхолдеры и стор ─────────────────────────────────────────────────

def test_extract_placeholders_ordered_unique():
    t = "Договор с {{Заказчик}} на {{сумма}}. Подпись: {{Заказчик}}"
    assert extract_placeholders(t) == ["Заказчик", "сумма"]


@pytest.fixture()
def store(tmp_path, monkeypatch):
    monkeypatch.setattr(fe, "_store_dir", lambda uid: tmp_path)
    return tmp_path


def test_template_and_requisites_store(store):
    rec = save_doc_template("u1", title="КП", body="Кому: {{Заказчик}}",
                            mode="strict")
    assert rec["placeholders"] == ["Заказчик"]
    assert fe.list_doc_templates("u1")[0]["id"] == rec["id"]
    fe.save_requisites("u1", {"ИНН": "7701234567", "Банк": "Т-Банк"})
    assert fe.get_requisites("u1")["ИНН"] == "7701234567"
    assert fe.delete_doc_template("u1", rec["id"]) is True


# ── strict-заполнение: реквизиты без LLM, деньги кодом, маркеры ────────

class _FakeLLM:
    async def generate_json(self, prompt=None, **kw):
        return {
            "fields": {
                "Заказчик": {"value": "ООО Ромашка",
                             "confidence": "found",
                             "source": "«договор с Ромашкой»"},
                "Срок": {"value": None},
            },
            "items": [{"name": "Аудит", "qty": 2, "unit": "шт",
                       "price": 100}],
        }


def test_fill_strict_end_to_end(store, monkeypatch):
    async def _no_meetings(uid, ids, cap_chars=0):
        return "=== ВСТРЕЧА: Продажи ===\nобсуждали договор"
    monkeypatch.setattr(fe, "_meetings_text", _no_meetings)
    fe.save_requisites("u1", {"ИНН": "7701234567"})

    res = asyncio.run(fe.fill_document(
        "u1",
        template_text=("Договор: {{Заказчик}}, ИНН исполнителя {{ИНН}}, "
                       "срок {{Срок}}. Итого: {{итого}} ({{сумма_прописью}}), "
                       "НДС {{ндс}}"),
        meeting_ids=["m1"], mode="strict", vat_rate=20, llm=_FakeLLM()))

    f = res["fields"]
    assert f["ИНН"]["confidence"] == CONF_FOUND        # из реквизитов, без LLM
    assert f["ИНН"]["source"] == "реквизиты компании"
    assert f["Заказчик"]["value"] == "ООО Ромашка"
    assert f["Срок"]["confidence"] == CONF_MISSING
    # деньги посчитал код: 2×100 +20% НДС
    assert f["итого"]["value"] == "240.00"
    assert "Двести сорок" in f["сумма_прописью"]["value"]
    assert f["ндс"]["value"] == "40.00"
    # документ: заполненное подставлено, missing остался {{Срок}}
    assert "ООО Ромашка" in res["document"]
    assert "{{Срок}}" in res["document"] or "{{ Срок }}" in res["document"]
    assert res["unfilled"] == ["Срок"]
    assert res["money"]["total"] == "240.00"


def test_fill_strict_llm_failure_graceful(store, monkeypatch):
    async def _mt(uid, ids, cap_chars=0):
        return "txt"
    monkeypatch.setattr(fe, "_meetings_text", _mt)

    class _Boom:
        async def generate_json(self, **kw):
            raise RuntimeError("llm down")
    res = asyncio.run(fe.fill_document(
        "u1", template_text="Поле: {{X}}", meeting_ids=[],
        mode="strict", llm=_Boom()))
    assert res["fields"]["X"]["confidence"] == CONF_MISSING
    assert res["unfilled"] == ["X"]


# ── docx-бланк: единый конвейер (сшивка с meeting_doc_engine) ───────────

def _make_docx_template() -> bytes:
    import io
    from docx import Document
    d = Document()
    d.add_paragraph("Договор с {{Заказчик}}")
    d.add_paragraph("ИНН исполнителя: {{ИНН}}")
    d.add_paragraph("Итого: {{итого}} ({{сумма_прописью}})")
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _docx_text(blob: bytes) -> str:
    import io
    import zipfile
    with zipfile.ZipFile(io.BytesIO(blob)) as z:
        return z.read("word/document.xml").decode("utf-8")


def test_fill_docx_template_end_to_end(store, monkeypatch):
    import base64

    async def _mt(uid, ids, cap_chars=0):
        return "=== ВСТРЕЧА ===\nобсуждали договор с Ромашкой"
    monkeypatch.setattr(fe, "_meetings_text", _mt)
    fe.save_requisites("u1", {"ИНН": "7701234567"})

    tpl = _make_docx_template()
    res = asyncio.run(fe.fill_document(
        "u1", meeting_ids=["m1"], template_docx=tpl,
        vat_rate=20, llm=_FakeLLM()))

    assert res["mode"] == "docx" and res["docx_b64"]
    xml = _docx_text(base64.b64decode(res["docx_b64"]))
    assert "ООО Ромашка" in xml            # из встречи (LLM)
    assert "7701234567" in xml             # реквизиты без LLM
    assert "240.00" in xml                 # деньги посчитал код (2×100 +НДС)
    assert "Двести сорок" in xml           # сумма прописью
    assert "{{Заказчик}}" not in xml
    # поле «Срок» в этом бланке отсутствует → unfilled пуст
    f = res["fields"]
    assert f["ИНН"]["source"] == "реквизиты компании"


def test_docx_template_store_roundtrip(store):
    import base64
    tpl = _make_docx_template()
    rec = fe.save_doc_template("u1", title="Бланк",
                               docx_b64=base64.b64encode(tpl).decode())
    assert rec["kind"] == "docx"
    assert "Заказчик" in rec["placeholders"]
    assert fe.load_template_docx("u1", rec["id"]) == tpl


def test_docx_engine_money_unified():
    """compute_line_items — адаптер над Decimal-движком (HALF_UP)."""
    from backend.core.documents.meeting_doc_engine import compute_line_items
    r = compute_line_items([{"name": "x", "qty": 1.5, "price": 99.99}],
                           vat_rate=0)
    assert r["lines"][0]["amount"] == 149.99   # HALF_UP (float-дубль давал .98)
    r2 = compute_line_items([{"name": "Работа", "qty": 2, "price": 100}],
                            vat_rate=20)
    assert r2["total"] == 240.0
    assert r2["total_in_words"] == "Двести сорок рублей 00 копеек"
